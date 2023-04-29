import requests
import pendulum
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Length
from docx.shared import RGBColor
from docx.table import _Cell


def insert_image(cell, screenshot_image, screenshot_folder):
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = 1  # center align the image

    # Save screenshots to fs
    screenshot_filename = screenshot_image.split('/')[-1]
    screenshot_file = requests.get(screenshot_image)
    open('{}/{}'.format(screenshot_folder, screenshot_filename), 'wb') \
        .write(screenshot_file.content)
    # insert image in cell
    cell_paragraph.add_run().add_picture(screenshot_folder + '/' + screenshot_filename,
                                         width=Inches(cell.width.inches))


def get_request(str_url, virtuoso_token):
    # print(str_url)
    headers = {"Authorization": "Bearer {}".format(virtuoso_token)}
    response = requests.get(str_url, headers=headers)
    if response.status_code != 200:
        raise Exception("Unexpected response code: {}".format(response.status_code))

    # print(response.json())
    if 'item' in response.json().keys():
        return response.json().get('item', {})
    if 'map' in response.json().keys():
        return response.json().get('map', {})
    else:
        return response.json()


def get_step_checkpoint_names(virtuoso_api, snapshot_id, goal_id, virtuoso_token):
    # get step details
    steps_details = get_request("https://{}/api/snapshots/{}/goals/{}/testsuites?envelope=false"
                                .format(virtuoso_api, snapshot_id, goal_id), virtuoso_token)

    all_steps = []
    checkpoint_title = {}
    for stepDetail in steps_details:
        cases = steps_details[stepDetail].get('cases', {})
        for case in cases:
            checkpoint_title[case.get('canonicalId')] = case.get('title')
            for s in case['steps']:
                all_steps.append(s)

    steps_text = {}
    steps_text_url = "https://step-deparser-service.virtuoso.workers.dev"
    steps_headers = {"Authorization": "Bearer {}".format(virtuoso_token)}
    steps_text_details = requests.post(steps_text_url, json={'steps': all_steps})

    j = 0
    for all_step in all_steps:
        steps_text[all_step['id']] = steps_text_details.json()[j]
        j = j + 1

    return checkpoint_title, steps_text


def add_space(space, doc):
    # Add a smaller space after the heading
    paragraph = doc.add_paragraph()
    paragraph.add_run('')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Length(Inches(0.04*space))  # 0.08 inch
    paragraph_format.space_after = Length(Inches(0.08*space))  # 0.04 inch


def format_duration(milliseconds):
    duration = pendulum.duration(milliseconds=milliseconds)
    return duration.in_words()


def outcome_color(outcome):
    if outcome == "PASS":
        return RGBColor(134, 184, 20)
    if outcome == "FAIL" or outcome == "ERROR":
        return RGBColor(203, 42,42)
    else:
        return RGBColor(206, 206, 206)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def format_table(tbl):
    tbl.style = 'Normal Table'
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"sz": 0, "val": "none", "color": "#FF0000", "space": "0"},
                bottom={"sz": 0, "color": "#00FF00", "val": "none", "space": "0"},
                start={"sz": 0, "val": "none", "shadow": "false", "space": "0"},
                end={"sz": 0, "val": "none", "space": "0"},
            )