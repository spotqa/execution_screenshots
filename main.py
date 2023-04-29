import os
import shutil
import docx
import datetime
import sys
import getopt
import time
import urllib.parse
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathvalidate import sanitize_filename
from functions import get_request, format_duration, get_step_checkpoint_names, outcome_color, insert_image, add_space


def check_args(argv):
    arg_token = ""
    arg_id = 0
    arg_env = "api-app2.virtuoso.qa"
    desc = "Virtuoso test execution report self extractor."
    arg_help = "\nUsage:  {0} OPTIONS\n\n{1}\n\nOptions:" \
               "\n\t-t, --token uuid    Virtuoso token" \
               "\n\t-i, --id int        Virtuoso execution id" \
               "\n\t-e, --env string    Environment [OPTIONAL] (default = \"api-app2.virtuoso.qa\")>" \
               "\n\nCopyright SPOTQA LTD 2023 | <support@virtuoso.qa>\n".format(argv[0], desc)

    def _help():
        print(arg_help) # print the help message
        sys.exit(2)

    try:
        opts, args = getopt.getopt(argv[1:], "ht:i:e:", ["help", "token=", "id=", "env="])
    except:
        _help()

    for opt, arg in opts:
        if opt in ("-h", "--help"):
            _help()
        elif opt in ("-t", "--token"):
            arg_token = arg
        elif opt in ("-i", "--id"):
            arg_id = arg
        elif opt in ("-e", "--env"):
            arg_env = arg

    if arg_token == "" or arg_id == "":
        _help()

    return arg_token, arg_id, arg_env


if __name__ == "__main__":
    VIRTUOSO_TOKEN, VIRTUOSO_EXECUTION_ID, VIRTUOSO_API = check_args(sys.argv)
    try :
        SCREENSHOT_FOLDER = "screenshots-{}".format(round(time.time() * 1000))
        VIRTUOSO_API_EXECUTION_DETAILS_PATH = "executions"
        DOC = docx.Document()
        print("Extracting report. This can take several minutes.")

        # Check if screenshots folder exists
        if not os.path.exists(SCREENSHOT_FOLDER):
            os.makedirs(SCREENSHOT_FOLDER)

        # Retrieve the screenshots from the execution details
        execDetails = get_request("https://{}/api/testsuites/execution?jobId={}&envelope=false"
                                  .format(VIRTUOSO_API, VIRTUOSO_EXECUTION_ID), VIRTUOSO_TOKEN)

        journeys = execDetails.get('journeys', {})

        allSteps = []
        screenshots = []
        init = False
        title = ""

        for journey in journeys:
            journeyDetail = journeys[journey].get('journey', {})
            snapshotId = journeyDetail['snapshotId']
            goalId = journeyDetail['goalId']

            # init doc title
            if not init:
                goalDetails = get_request("https://{}/api/goals/{}?goalSnapshotId={}&envelope=false"
                                          .format(VIRTUOSO_API, goalId, snapshotId), VIRTUOSO_TOKEN)

                title = goalDetails.get('name') + ' Execution Report'
                DOC.add_heading(title, 0)

                # get submit date
                submitDate = journeys[journey].get('lastExecution', {}).get('job', {}).get('submitDate')
                # Convert the timestamp to a datetime object
                dt = datetime.datetime.fromtimestamp(submitDate / 1000.0)
                # Format the datetime object as a string
                dateString = dt.strftime('%A, %d %B %Y')

                # get journey statistics
                stats = execDetails.get('journeyStatistics', {})

                DOC.add_paragraph()\
                    .add_run('Submitted on {}. Completed with "{}" status, in {}'
                             .format(dateString, stats.get('outcome'), format_duration(stats.get('totalDuration'))))
                init = True

            checkpointTitle, stepsText = get_step_checkpoint_names(VIRTUOSO_API, snapshotId, goalId, VIRTUOSO_TOKEN)

            checkpoints = journeys[journey].get('lastExecution', {}).get('report', {}).get('checkpoints', {})

            DOC.add_heading(journeys[journey].get('journey').get('title'), 1)

            for checkpoint in checkpoints:
                DOC.add_heading(checkpointTitle[checkpoint], 2)
                steps = checkpoints[checkpoint].get('steps', {})

                for step in steps:
                    stepId = steps[step].get('stepId')

                    heading = DOC.add_heading("", 3)
                    # add a run to the heading and set its font color to red
                    heading_run = heading.add_run(stepsText[stepId])
                    heading_run.font.color.rgb = outcome_color(steps[step].get('outcome'))  # red color

                    beforeScreenshot = steps[step].get('beforeScreenshot')
                    screenshot = steps[step].get('screenshot')

                    if beforeScreenshot or screenshot:
                        table = DOC.add_table(rows=2, cols=2 if beforeScreenshot and screenshot else 1)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    if beforeScreenshot:
                        cell1 = table.rows[0].cells[0]
                        cell2 = table.rows[1].cells[0]
                        cell1_paragraph = cell1.paragraphs[0]
                        cell1_paragraph.alignment = 0  # center align the image
                        cell1_paragraph.add_run('Screenshot before step execution')
                        insert_image(cell2, beforeScreenshot, SCREENSHOT_FOLDER)

                    if screenshot:
                        cell1 = table.rows[0].cells[1 if beforeScreenshot else 0]
                        cell2 = table.rows[1].cells[1 if beforeScreenshot else 0]
                        cell1_paragraph = cell1.paragraphs[0]
                        cell1_paragraph.alignment = 0  # center align the image
                        cell1_paragraph.add_run('Screenshot after step execution')
                        insert_image(cell2, screenshot, SCREENSHOT_FOLDER)

                    add_space(2, DOC)

        fname = sanitize_filename(title + '.docx')
        DOC.save(fname)
        if os.path.exists(SCREENSHOT_FOLDER):
            shutil.rmtree(SCREENSHOT_FOLDER)
        print("Report extracted to: file://{}".format(urllib.parse.quote(os.path.abspath(fname))))
    except Exception as e:
        if os.path.exists(SCREENSHOT_FOLDER):
            shutil.rmtree(SCREENSHOT_FOLDER)
        sys.exit("Report extraction failed with the following error:\n\t{}".format(e))
